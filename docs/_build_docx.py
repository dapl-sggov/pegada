"""Gera 14a_Questionario_OnePager.docx a partir do conteudo estruturado.

Identidade visual: Calibri/Cambria como fallback (DOCX usa fontes do sistema).
Cores: azul governo #0a3161 + dourado #b08020.
"""
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


GOV_BLUE = RGBColor(0x0A, 0x31, 0x61)
GOV_BLUE_DARK = RGBColor(0x06, 0x20, 0x3F)
GOLD = RGBColor(0xB0, 0x80, 0x20)
INK = RGBColor(0x0C, 0x17, 0x29)
INK_MUTED = RGBColor(0x5B, 0x64, 0x78)
INK_FAINT = RGBColor(0x9A, 0xA5, 0xB6)
BG_SOFT = RGBColor(0xF1, 0xF4, 0xF8)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    """Aplica bordas — kwargs: top, bottom, start, end com {'sz':n,'color':'XXXXXX'}."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side, props in kwargs.items():
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(props.get('sz', 4)))
        border.set(qn('w:color'), props.get('color', '000000'))
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def style_run(run, *, font='Calibri', size=11, bold=False, color=None, italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # Garante font fallback em ambiente Windows
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font)
    r_fonts.set(qn('w:hAnsi'), font)
    r_fonts.set(qn('w:cs'), font)


def add_para(doc, text, *, font='Calibri', size=11, bold=False, color=None, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    style_run(run, font=font, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading_band(doc, text, *, size=22, color=GOV_BLUE_DARK):
    # Faixa de título com fundo azul claro
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    style_run(run, font='Cambria', size=size, bold=True, color=color)
    return p


def add_eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    style_run(run, font='Calibri', size=8, bold=True, color=GOLD)
    # Letter-spacing simulado por espaços não disponível; pequeno OK.
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    style_run(run, font='Cambria', size=13, bold=True, color=GOV_BLUE_DARK)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    style_run(run, font='Cambria', size=11, bold=True, color=GOV_BLUE)
    return p


def add_rich(doc, parts, *, size=10.5, space_after=4):
    """Parágrafo multi-run. parts: lista de (texto, opts)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, opts in parts:
        run = p.add_run(text)
        style_run(run, size=size, **opts)
    return p


def add_horizontal_rule(doc, color='D9E0EA'):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build():
    doc = Document()
    # Margens institucionais
    for section in doc.sections:
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)

    # Cabeçalho — eyebrow + título
    add_eyebrow(doc, "DAPL/DSSD · Pegada Legislativa · Maio 2026")
    add_heading_band(doc, "Pegada Legislativa — pré-instalação na RING", size=20)
    add_para(doc, "Agenda de uma página · DAPL/DSSD → DSTD + SmartLegis",
             font='Calibri', size=10, italic=True, color=INK_MUTED, space_after=10)

    add_horizontal_rule(doc)

    # Contexto
    add_h2(doc, "Contexto em três frases")
    add_rich(doc, [
        ("A ", {}),
        ("Lei n.º 5-A/2026", {'bold': True, 'color': INK}),
        (" obriga a documentar interlocutores e contributos no processo legislativo. Para tornar esta obrigação efetiva, a ", {}),
        ("DAPL/DSSD", {'bold': True, 'color': INK}),
        (" desenhou um sistema de ", {}),
        ("quatro marcos bloqueantes (M0/M1/M4/M5)", {'bold': True, 'color': INK}),
        (" — com a consulta pública a ocorrer entre M1 (Pré-RSE) e M4 (Pré-CM) — cujo cumprimento é verificado pelo ", {}),
        ("SmartLegis", {'bold': True, 'color': INK}),
        (" antes de autorizar a tramitação para RSE, CM e DR. O acoplamento é assíncrono e criptográfico: a FPL Ponte emite um ", {}),
        ("JWS Ed25519", {'bold': True, 'color': INK}),
        (" em cada marco; o SmartLegis verifica-o offline com a chave pública previamente entregue. As restantes integrações (diretório, SMTP, RTRI, ConsultaLEX, DRE) são não-bloqueantes com fallback manual.", {}),
    ], size=10.5, space_after=8)

    add_rich(doc, [
        ("Documento completo: ", {'color': INK_MUTED}),
        ("docs/14_Questionario_Infraestrutura.md", {'color': GOV_BLUE, 'font': 'Consolas'}),
    ], size=9.5, space_after=10)

    # O que já entregamos
    add_h2(doc, "O que o nosso lado já entrega")
    add_para(doc, "Imagem Docker + docker-compose (Postgres/Redis/MinIO) · CSP/CSRF/2FA TOTP · spec do comprovativo fechada (docs/12) · adapters mock/manual para integrações · runbook (docs/06) · DPIA + threat models · análise NIS2 da aplicação (docs/15).",
             size=10, color=INK, space_after=10)

    add_horizontal_rule(doc)

    # Decisões bloqueantes - tabela
    add_h2(doc, "Decisões bloqueantes")
    add_para(doc, "Precisam de resposta antes da geração da chave Ed25519.",
             size=9.5, italic=True, color=INK_MUTED, space_after=6)

    table = doc.add_table(rows=7, cols=4)
    table.autofit = False
    col_widths = [Cm(0.8), Cm(5.8), Cm(2.8), Cm(6.6)]
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = w

    # Header
    hdr_cells = table.rows[0].cells
    hdr_data = [("#", ""), ("Decisão", ""), ("Quem decide", ""), ("Porquê é bloqueante", "")]
    for cell, (text, _) in zip(hdr_cells, hdr_data):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, "0A3161")
        set_cell_border(cell, bottom={'sz': 12, 'color': 'B08020'})
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        style_run(run, font='Calibri', size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    decisoes = [
        ("1", "Nome DNS interno definitivo (sugestão fpl.gov.pt)", "DSTD + DAPL/DSSD", "Hardcoded no iss do JWS — mudar implica nova chave e nova distribuição"),
        ("2", "Cofre de segredos e modo da chave privada Ed25519 (HSM ou ficheiro 0600?)", "DSTD", "Afeta o módulo de assinatura"),
        ("3", "Calendário do SmartLegis para o verificador (T0 + 10 sem é viável até 27 jul?)", "SmartLegis", "Prazo legal de entrada em vigor"),
        ("4", "Diretório interno — LDAP/AD URL, base DN, conta de bind, mapeamento de grupos", "DSTD", "Sem isto fica em modo local (não escala)"),
        ("5", "VLAN, firewall ingress/egress e reverse proxy interno", "DSTD", "Sem ingress não há acesso; sem egress integrações ficam em manual"),
        ("6", "Validação NIS2 institucional (CISO assina configuração + canal de notificação CNCS)", "DSTD", "Sem isto, o go-live não tem cobertura institucional NIS2"),
    ]
    for row_i, dec in enumerate(decisoes, start=1):
        row_cells = table.rows[row_i].cells
        for col_i, text in enumerate(dec):
            cell = row_cells[col_i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if row_i % 2 == 0:
                set_cell_bg(cell, "F6F8FB")
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            run = para.add_run(text)
            bold = (col_i == 0)
            color = GOV_BLUE if col_i == 0 else INK
            size = 9.5
            style_run(run, font='Calibri', size=size, bold=bold, color=color)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_horizontal_rule(doc)

    # Outras perguntas
    add_h2(doc, "Outras perguntas críticas (não bloqueantes mas urgentes)")
    perguntas = [
        ("Serviços de dados", "Postgres 16 / Redis 7 / MinIO geridos pela DSTD ou container nosso?"),
        ("SMTP do Estado", "Host, porta, endereço From, limites de throughput?"),
        ("RTRI / DRE", "Acordos de uso das APIs? Calendário? (ConsultaLEX agora coordenado internamente com a DSTD.)"),
        ("Distribuição da imagem", "GHCR → registry interno? Procedimento de promoção?"),
        ("Backups e retenção", "RPO/RTO? Storage WORM para arquivo de 10 anos dos comprovativos?"),
        ("Monitorização", "Prometheus centralizado? AlertManager? Agregação de logs?"),
        ("Operação 24/7", "DSTD, DAPL/DSSD em horário útil, ou modelo misto?"),
        ("Conformidade NIS2", "Análise da app face ao art. 21.º já em docs/15 (5 cumprem · 3 parciais · 1 a iniciar · 1 dependente · nenhum em incumprimento). Falta integração institucional: CISO, SOC/SIEM, notificação CNCS, encriptação em repouso."),
    ]
    for label, q in perguntas:
        add_rich(doc, [
            (label + ": ", {'bold': True, 'color': GOV_BLUE_DARK}),
            (q, {'color': INK}),
        ], size=10, space_after=3)

    add_horizontal_rule(doc)

    # SmartLegis
    add_h2(doc, "O que o SmartLegis tem de desenvolver")
    sl = [
        "Verificador JWS Ed25519 (~6 semanas) — algoritmo em docs/12 §4. Bibliotecas standard, sem dependência síncrona da FPL Ponte.",
        "Regra de aceitação por estado — RSE exige M0+M1; CM exige +M4; DR exige +M5. (CP entre RSE e CM: M2 abre, M3 encerra; ambos informativos.)",
        "Bloqueio com mensagem específica quando falha (alg-recusado, kid-desconhecido, assinatura-invalida, expirado, ...).",
        "Validação contra vetores de teste TV1–TV6 antes do go-live.",
    ]
    for i, item in enumerate(sl, start=1):
        add_rich(doc, [
            (f"  {i}. ", {'bold': True, 'color': GOLD}),
            (item, {'color': INK}),
        ], size=10, space_after=3)

    add_horizontal_rule(doc)

    # Pedido da reunião
    add_h2(doc, "Pedido concreto desta reunião")
    pedidos = [
        "Calendarizar resposta às 6 decisões bloqueantes (proposta: 10 dias úteis).",
        "Identificar o ponto de contacto único na DSTD para esta instalação.",
        "Confirmar o calendário do SmartLegis ou identificar dependências que o ponham em risco.",
        "Marcar reunião técnica de seguimento com a especificação do comprovativo (docs/12) aberta para Q&A.",
    ]
    for i, item in enumerate(pedidos, start=1):
        add_rich(doc, [
            (f"  {i}. ", {'bold': True, 'color': GOV_BLUE}),
            (item, {'color': INK}),
        ], size=10, space_after=3)

    # Prazo legal — destaque
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    box = doc.add_table(rows=1, cols=1)
    cell = box.rows[0].cells[0]
    set_cell_bg(cell, "FFF8E6")
    set_cell_border(cell, top={'sz': 12, 'color': 'B08020'}, bottom={'sz': 12, 'color': 'B08020'},
                    start={'sz': 12, 'color': 'B08020'}, end={'sz': 12, 'color': 'B08020'})
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("PRAZO LEGAL  ·  ")
    style_run(run, size=10, bold=True, color=GOLD)
    run = p.add_run("27 de julho de 2026")
    style_run(run, size=12, bold=True, color=GOV_BLUE_DARK)
    run = p.add_run("  — entrada em vigor da Lei n.º 5-A/2026. Caminho crítico: instalação na RING + verificador no SmartLegis + ligação ao diretório + DPIA submetida à CNPD.")
    style_run(run, size=9.5, color=INK)

    # Footer
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_para(doc, "Documento completo, perguntas detalhadas com impacto por decisão, e checklist de go-live em docs/14_Questionario_Infraestrutura.md.",
             size=8.5, italic=True, color=INK_FAINT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '14a_Questionario_OnePager.docx')
    doc.save(out)
    print(f"OK: {out}")


if __name__ == '__main__':
    build()
